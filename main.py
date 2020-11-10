import docx
import os

def Insert_Image_Document(document,image_path):
    document.add_picture(image_path,width=docx.shared.Cm(15),height=docx.shared.Cm(10))

def main():
    images = os.listdir("screenshots")
    #init document
    document = docx.Document()
    document.add_heading('Work', 0)
    #go through image path list
    for image in images:
        Insert_Image_Document(document,"screenshots/"+image)
    document.save("document.docx")

main()